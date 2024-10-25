import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx2pdf import convert as docx_to_pdf
from pdf2docx import parse as pdf_to_docx
import pandas as pd
from pptx import Presentation
from fpdf import FPDF
from docx import Document
import fitz  # PyMuPDF
import shutil


class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Dosya Dönüştürme Uygulaması")
        self.root.geometry("600x400")

        # Hoşgeldiniz mesajı
        tk.Label(root, text="Hoşgeldiniz!", font=("Helvetica", 16)).pack(pady=10)

        # Dönüştürme talebi
        tk.Label(root, text="Neye dönüştürmek istersiniz?", font=("Helvetica", 12)).pack(pady=5)

        # Çıktı formatı seçimi
        self.output_format = tk.StringVar(value="Seçiniz")
        formats = ["PDF", "Word", "Excel", "CSV", "Metin"]
        tk.OptionMenu(root, self.output_format, *formats).pack(pady=5)

        # Dosya seçme butonu
        tk.Button(root, text="Dosya Seç", command=self.select_file).pack(pady=10)

        # Bilgilendirme label
        self.info_label = tk.Label(root, text="", fg="green")
        self.info_label.pack(pady=5)

    def select_file(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            input_ext = os.path.splitext(file_path)[1].lower()
            output_format = self.output_format.get()
            output_ext = self.get_output_extension(output_format)

            if output_ext is None:
                messagebox.showwarning("Uyarı", "Lütfen geçerli bir çıktı formatı seçin.")
                return

            # Dönüştürme kontrolü
            if self.can_convert(input_ext, output_ext):
                self.convert_file(file_path, output_ext)
            else:
                messagebox.showerror("Hata", f"{input_ext} dosyası {output_ext} formatına dönüştürülemiyor.")

    def get_output_extension(self, format_str):
        format_mapping = {
            "PDF": "pdf",
            "Word": "docx",
            "Excel": "xlsx",
            "CSV": "csv",
            "Metin": "txt"
        }
        return format_mapping.get(format_str)

    def can_convert(self, input_ext, output_ext):
        # Belirli dönüşüm kombinasyonları
        conversion_map = {
            '.docx': ['pdf', 'docx', 'txt'],
            '.pdf': ['docx', 'txt'],
            '.xlsx': ['csv', 'xlsx'],
            '.csv': ['xlsx', 'csv'],
            '.pptx': ['pdf'],
            '.txt': ['pdf', 'docx', 'txt']
        }
        return output_ext in conversion_map.get(input_ext, [])

    def convert_file(self, input_path, output_ext):
        try:
            output_dir = os.path.dirname(input_path)
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_file = os.path.join(output_dir, f"{base_name}.{output_ext}")
            input_ext = os.path.splitext(input_path)[1].lower()

            # Dönüştürme işlemleri
            converters = {
                ('.docx', 'pdf'): lambda: docx_to_pdf(input_path, output_file),
                ('.docx', 'docx'): lambda: shutil.copy(input_path, output_file),
                ('.docx', 'txt'): lambda: self.docx_to_txt(input_path, output_file),
                ('.pdf', 'docx'): lambda: self.pdf_to_docx(input_path, output_file),
                ('.pdf', 'txt'): lambda: self.pdf_to_txt(input_path, output_file),
                ('.xlsx', 'csv'): lambda: self.xlsx_to_csv(input_path, output_file),
                ('.csv', 'xlsx'): lambda: self.csv_to_xlsx(input_path, output_file),
                ('.pptx', 'pdf'): lambda: self.pptx_to_pdf(input_path, output_file),
                ('.txt', 'pdf'): lambda: self.txt_to_pdf(input_path, output_file),
                ('.txt', 'docx'): lambda: self.txt_to_docx(input_path, output_file),
                ('.txt', 'txt'): lambda: shutil.copy(input_path, output_file)
            }

            convert = converters.get((input_ext, output_ext))
            if convert:
                convert()
                self.info_label.config(text=f"{output_ext.upper()} dosyası {output_dir} klasörüne kaydedildi.")
            else:
                messagebox.showerror("Hata", f"{input_ext} dosyası {output_ext} formatına dönüştürülemiyor.")

        except Exception as e:
            messagebox.showerror("Dönüştürme Hatası", f"Dönüştürme işlemi başarısız oldu.\nHata: {e}")

    def docx_to_txt(self, input_path, output_path):
        doc = Document(input_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))

    def pdf_to_docx(self, input_path, output_path):
        pdf_to_docx(pdf_file=input_path, docx_file=output_path, start=0, end=None)

    def pdf_to_txt(self, input_path, output_path):
        with fitz.open(input_path) as doc:
            text = ""
            for page in doc:
                text += page.get_text()
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)

    def xlsx_to_csv(self, input_path, output_path):
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False)

    def csv_to_xlsx(self, input_path, output_path):
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False)

    def pptx_to_pdf(self, input_path, output_path):
        from comtypes import client

        powerpoint = client.CreateObject('Powerpoint.Application')
        powerpoint.Visible = 1

        deck = powerpoint.Presentations.Open(input_path)
        deck.SaveAs(output_path, 32)  # 32 for ppt to pdf
        deck.Close()
        powerpoint.Quit()

    def txt_to_pdf(self, input_path, output_path):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(input_path, 'r', encoding='utf-8') as f:
            for line in f:
                pdf.multi_cell(0, 10, line)
        pdf.output(output_path)

    def txt_to_docx(self, input_path, output_path):
        doc = Document()
        with open(input_path, 'r', encoding='utf-8') as f:
            doc.add_paragraph(f.read())
        doc.save(output_path)


if __name__ == "__main__":
    root = tk.Tk()
    app = FileConverterApp(root)
    root.mainloop()
